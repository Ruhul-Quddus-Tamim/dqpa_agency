from agency_swarm.tools import BaseTool
from pydantic import Field
from docx import Document


class ReportGenerationTool(BaseTool):
    """
    Generates comprehensive reports in document format incorporating findings and suggestions for data quality improvements.
    """

    findings: str = Field(
        ..., description="String containing the findings from the data quality checks."
    )
    suggestions: str = Field(
        ..., description="String containing suggestions for improving data quality."
    )
    report_title: str = Field(
        ..., description="Title of the report."
    )

    def run(self):
        doc = Document()
        doc.add_heading(self.report_title, 0)
        doc.add_heading('Findings', level=1)
        doc.add_paragraph(self.findings)
        doc.add_heading('Suggestions', level=1)
        doc.add_paragraph(self.suggestions)
        # More sections can be added as needed

        # Save the document
        doc.save(f'{self.report_title}.docx')

        return f"Report '{self.report_title}' generated successfully."
